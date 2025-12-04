from docx import Document
from docx.table import Table, _Row, _Cell
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from typing import List, Dict, Any
from pathlib import Path
from copy import deepcopy
import re
import logging



logger = logging.getLogger(__name__)

class Word:
    
    def __init__(self, template_path: str):
        self.word = Document(template_path)
        self.placeholders = set()
    
    def find_all_placeholders(self, pattern: str = r'\{\{(\w+)\}\}') -> List[str]:
        ph = set()
        
        #Placeholder Paragrafos
        for p in self.word.paragraphs:
            ph.update(re.findall(pattern, p.text))
        
        #Placeholder Tabela
        for t in self.word.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        ph.update(re.findall(pattern, p.text))
        
        for s in self.word.sections:
            #Placeholder Cabeçalho
            for h in s.header.paragraphs:
                ph.update(re.findall(pattern, h.text))
            #Placeholder Rodapé
            for f in s.footer.paragraphs:
                ph.update(re.findall(pattern, f.text))
        
        self.placeholders = ph
        logger.info(f"Placeholders encontrados: {self.placeholders}")
        return list(ph)
    
    #contabilizar quantas linhas na tabela tem placeholder
    def f_table_ph(self, table: Table, pattern: str = r'\{\{(\w+)\}\}'):
        ph = set()
        for r in table.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    ph.update(re.findall(pattern, p.text))
        f = 1 + len(ph)
        return f
    
    def insert_rowP(self, table: Table, i):
        new_row = parse_xml(f'<w:tr {nsdecls("w")}/>')
        target_row = table.rows[i]._element
        target_row.addprevious(new_row)
        new = _Row(new_row, table)
        return new
    
    def insert_row(self, table: Table, index: int) -> _Row:
        # Pega uma linha de referência para copiar estrutura
        if index > 0 and index <= len(table.rows):
            reference_row = table.rows[index - 1]
        else:
            reference_row = table.rows[0]
        
        num_cols = len(reference_row.cells)
        logger.info(f"[INSERT_ROW] Criando linha no indice {index} com {num_cols} celulas")
        
        # Cria novo elemento de linha
        new_row_element = parse_xml(f'<w:tr {nsdecls("w")}/>')
        
        # Adiciona células à nova linha baseado na linha de referência
        for i in range(num_cols):
            ref_cell = reference_row.cells[i]._element
            
            # Cria célula vazia
            cell_xml = parse_xml(f'<w:tc {nsdecls("w")}><w:p><w:r><w:t></w:t></w:r></w:p></w:tc>')
            
            new_row_element.append(cell_xml)
            
            # Opcional: Copia propriedades da célula de referência (largura, bordas, etc)
            # Esta dando problema na formatação!
            '''ref_cell = reference_row.cells[i]._element
            new_cell = cell_xml
            
            # Copia propriedades de célula (tcPr)
            if ref_cell.tcPr is not None:
                new_cell.tcPr = ref_cell.tcPr.__copy__()
                #tcPr_copy = deepcopy(ref_cell.tcPr)
                #cell_xml.insert(0, tcPr_copy)
                '''
                
        
        # Pega a linha alvo
        target_row = table.rows[index]._element
        
        # Insere antes da linha alvo
        target_row.addprevious(new_row_element)
        
        # Cria objeto Row do python-docx
        new_row = _Row(new_row_element, table)
        
        logger.info(f"[INSERT_ROW] Linha criada com sucesso: {len(new_row.cells)} celulas")
        
        return new_row
    
    
    def wanted_data(self, data: List[Dict[str,Any]], ano, bimestre):
        d = {}
        l = []
        l.append(data['values'][0])
        
        # Cuidar com °
        for v in range(len(data['values'])):
            if data['values'][v][1] == ano and data['values'][v][2] == bimestre:
                l.append(data['values'][v])
        d.update({'values': l})
        return d
    
    def mapa_tabela(self, data: List[Dict[str,Any]]):
        d= {}
        for h in data['values'][0]:
            d.update({h:[]})
        db = list(map(list,zip(*data['values'])))
        for l in range(len(db)):
            dados = []
            for v in range(len(db[l])):
                if v == 0:
                    continue
                dados.append(db[l][v])
            d.update({db[l][0]:dados})
        return d
    
    def fill_table_GA(self, data: List[Dict[str, Any]], ano, bimestre):
        if not data:
            logger.info("Sem Dados para o Preenchimento")
            return
        
        PROTECTED_ROWS = 6 # Número de linhas protegidas no final
        HEADER_OFFSET = 8 # Índice onde começam os dados (após cabeçalho)
        
        for table in self.word.tables:
            l = self.f_table_ph(table)
            
            header_row:_Row = table.rows[7]
            headers = [cell.text.strip().upper() for cell in header_row.cells]
            col_map = {}
            logger.info(f"Headers: {headers}")
            
            for i, header in enumerate(headers):
                if 'TÍTULO' in header or 'TITULO' in header:
                    col_map['TÍTULO'] = i
                elif 'CONTEÚDO' in header:
                    col_map['CONTEÚDOS'] = i
                elif 'HABILIDADES' in header:
                    col_map['HABILIDADES'] = i
                elif 'OBJETIVO' in header:
                    col_map['OBJETIVOS'] = i

            logger.info(f"Mapeamento: {col_map}")
            
            new_data = self.wanted_data(data, ano, bimestre)
            dados = self.mapa_tabela(new_data)
            logger.info(dados)
            index = len(dados['Bimestre'])
            
            logger.info(f"Dados: {dados}")
            logger.info(f"Total de linhas de dados: {index}")
            
            total_rows = len(table.rows)
            logger.info(f"Total de linhas na tabela: {total_rows}")
            
            # Área editável: de HEADER_OFFSET até (total_rows - PROTECTED_ROWS)
            first_editable_row = HEADER_OFFSET
            last_editable_row = total_rows - PROTECTED_ROWS
            max_data_rows = last_editable_row - first_editable_row

            logger.info(f"Area editavel: linhas {first_editable_row} a {last_editable_row-1}")
            logger.info(f"Linhas protegidas: {last_editable_row} a {total_rows-1}")
            logger.info(f"Maximo de linhas de dados: {max_data_rows}")
            
            """data_rows = len(table.rows) - l - 4
            
            logger.info(f"Dados: {dados}")
            logger.info(f"Index: {index}, data_rows: {data_rows}")
            
            rows_needed = index + 5
            current_rows = len(table.rows) - 8  # Linhas disponíveis após cabeçalho
        
            logger.info(f"Linhas necessarias: {rows_needed}, Linhas atuais: {current_rows}")
            
            # Adiciona linhas faltantes
            if rows_needed > current_rows:
                for _ in range(rows_needed - current_rows):
                    table.add_row()
                    logger.info(f"Linha adicionada. Total agora: {len(table.rows)}")
            """
            
            # VERIFICA se precisa adicionar linhas
            if index > max_data_rows:
                rows_to_add = index - max_data_rows
                logger.info(f"Necessario adicionar {rows_to_add} linhas")
                
                # Adiciona linhas ANTES das linhas protegidas
                # Insere na posição last_editable_row (antes das protegidas)
                for _ in range(rows_to_add):
                    # IMPORTANTE: Usa insert_row para adicionar ANTES das linhas protegidas
                    self.insert_row(table, last_editable_row)
                    last_editable_row += 1  # Atualiza posição
                    logger.info(f"Linha adicionada. Novo limite: {last_editable_row}")
            else:
                logger.info(f"Linhas suficientes. Usando {index} de {max_data_rows} disponiveis")
            
            
            # i pode ir nos valores esepcificos por bimestre/ano
            for i in range(index):
                ri = i + HEADER_OFFSET # Índice da linha na tabela
                logger.info(f"\n--- Processando linha de dados {i+1}/{index} (indice tabela: {ri}) ---")
                
                # VALIDAÇÃO: Não permite preencher linhas protegidas
                if ri >= (len(table.rows) - PROTECTED_ROWS):
                    logger.error(f"ERRO: Tentativa de sobrescrever linha protegida {ri}!")
                    logger.error(f"Parando preenchimento para proteger dados.")
                    break
                
                # Usa linha existente
                if ri < len(table.rows):
                    row = table.rows[ri]
                    logger.info(f"Usando linha existente {ri}: {len(row.cells)} celulas")
                else:
                    logger.error(f"ERRO: Linha {ri} nao existe! Total de linhas: {len(table.rows)}")
                    continue
                
                # Preenche células
                for field, col_idx in col_map.items():
                    if col_idx is None:
                        logger.warning(f"PROBLEMA: col_idx é None para campo '{field}'")
                        continue
                
                    if col_idx >= len(row.cells):
                        logger.warning(f"PROBLEMA: col_idx ({col_idx}) >= len(row.cells) ({len(row.cells)})")
                        continue
                    
                    logger.info(f"Condição OK! Preenchendo célula [{ri},{col_idx}]")
                    
                    cell:_Cell = row.cells[col_idx]
                    
                    value = None
                    tentativas = [
                        field,                   
                        field.upper(),            
                        field.capitalize(),       
                        field.title(),            
                    ]

                    for tentativa in tentativas:
                        if tentativa in dados.keys():
                            value = dados[tentativa][i]
                            logger.info(f"Valor encontrado com chave '{tentativa}': '{value}'")
                            break
                    else:
                        logger.warning(f"Valor não encontrado. Tentativas: {tentativas}")
                        logger.warning(f"Chaves disponíveis: {list(dados.keys())}")
                        #value = dados.get(field)
                        #logger.info(value)
                    
                    #cell.add_paragraph(str(value[i]) if value is not None and value != '' else '')
                    texto = str(value) if value is not None and value != '' else ''

                    # Limpa célula e adiciona novo texto
                    cell.text = ''  # Limpa conteúdo existente
                    if texto:
                        cell.paragraphs[0].text = texto  # Usa o parágrafo existente
                    logger.info(f"Celula preenchida: '{texto}'")
            logger.info(f"\n[SUCESSO] Preenchimento concluido. Linhas protegidas preservadas.")
    # função para a troca de placeholders por conteudo
    def replace_placeholders(self, dados: dict, pattern: str = r'\{\{(\w+)\}\}'):
        for t in self.word.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        for key, value in dados.items():
                            placeholder = "{"+"{"+key+"}"+"}"
                            if pattern in p.text:
                                inline = p.runs
                                for item in inline:
                                    if key in item.text:
                                        item.text = item.text.replace(placeholder, value) 
    
    def save_document(self, output_path: str):
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.word.save(output_path)
        logger.info(f"Documento salvo: {output_path}")